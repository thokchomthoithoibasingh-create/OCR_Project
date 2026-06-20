"""
output_writer.py
-----------------
Takes raw OCR results and saves to both .txt and .docx files.
"""

import os
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def format_results_for_display(results: list, show_confidence: bool = True) -> str:
    if not results:
        return "[No text detected]"
    lines = []
    for item in results:
        if show_confidence:
            conf_pct = item["confidence"] * 100
            lines.append(f"{item['text']}    [confidence: {conf_pct:.2f}%]")
        else:
            lines.append(item["text"])
    return "\n".join(lines)


def format_plain_text(results: list) -> str:
    if not results:
        return ""
    return "\n".join(item["text"] for item in results)


def calculate_average_confidence(results: list) -> float:
    if not results:
        return 0.0
    total = sum(item["confidence"] for item in results)
    return round((total / len(results)) * 100, 2)


def _build_image_docx(results: list, source_filename: str):
    doc = Document()
    doc.add_heading("OCR Output Report", level=1)
    info = doc.add_paragraph()
    info.add_run(f"Source File       : {os.path.basename(source_filename)}\n")
    info.add_run(f"Processed On      : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    info.add_run(f"Lines Detected    : {len(results)}\n")
    info.add_run(f"Average Confidence: {calculate_average_confidence(results)}%\n")
    doc.add_paragraph("─" * 40)
    doc.add_heading("Extracted Text", level=2)
    for item in results:
        doc.add_paragraph(item["text"])
    return doc


def _build_pdf_docx(page_results: list, source_filename: str):
    doc = Document()
    doc.add_heading("OCR Output Report (Multi-Page PDF)", level=1)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    total_lines = sum(len(p) for p in page_results)
    all_conf = [item["confidence"] for page in page_results for item in page]
    avg_conf = round((sum(all_conf) / len(all_conf)) * 100, 2) if all_conf else 0.0
    info = doc.add_paragraph()
    info.add_run(f"Source File       : {os.path.basename(source_filename)}\n")
    info.add_run(f"Processed On      : {timestamp}\n")
    info.add_run(f"Total Pages       : {len(page_results)}\n")
    info.add_run(f"Total Lines       : {total_lines}\n")
    info.add_run(f"Average Confidence: {avg_conf}%\n")
    for page_num, results in enumerate(page_results, start=1):
        doc.add_paragraph("─" * 40)
        page_avg = calculate_average_confidence(results)
        doc.add_heading(f"Page {page_num}  (Lines: {len(results)}, Avg Confidence: {page_avg}%)", level=2)
        if results:
            for item in results:
                doc.add_paragraph(item["text"])
        else:
            doc.add_paragraph("[No text detected on this page]")
    return doc


def save_image_ocr_output(results: list, source_filename: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(source_filename))[0]

    # TXT
    txt_path = os.path.join(output_dir, f"{base_name}_ocr_output.txt")
    avg_conf = calculate_average_confidence(results)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header = (
        f"OCR Output Report\n"
        f"{'=' * 50}\n"
        f"Source File       : {os.path.basename(source_filename)}\n"
        f"Processed On      : {timestamp}\n"
        f"Lines Detected    : {len(results)}\n"
        f"Average Confidence: {avg_conf}%\n"
        f"{'=' * 50}\n\n"
    )
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(header)
        f.write(format_plain_text(results))
        f.write("\n")

    # DOCX
    docx_path = None
    if DOCX_AVAILABLE:
        docx_path = os.path.join(output_dir, f"{base_name}_ocr_output.docx")
        _build_image_docx(results, source_filename).save(docx_path)

    return txt_path, docx_path


def save_pdf_ocr_output(page_results: list, source_filename: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(source_filename))[0]

    # TXT
    txt_path = os.path.join(output_dir, f"{base_name}_ocr_output.txt")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    total_lines = sum(len(p) for p in page_results)
    all_conf = [item["confidence"] for page in page_results for item in page]
    avg_conf = round((sum(all_conf) / len(all_conf)) * 100, 2) if all_conf else 0.0
    header = (
        f"OCR Output Report (Multi-Page PDF)\n"
        f"{'=' * 50}\n"
        f"Source File       : {os.path.basename(source_filename)}\n"
        f"Processed On      : {timestamp}\n"
        f"Total Pages       : {len(page_results)}\n"
        f"Total Lines       : {total_lines}\n"
        f"Average Confidence: {avg_conf}%\n"
        f"{'=' * 50}\n\n"
    )
    body_parts = []
    for page_num, results in enumerate(page_results, start=1):
        page_avg = calculate_average_confidence(results)
        page_header = f"\n--- Page {page_num} (Lines: {len(results)}, Avg Confidence: {page_avg}%) ---\n"
        page_text = format_plain_text(results) if results else "[No text detected on this page]"
        body_parts.append(page_header + page_text)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(header)
        f.write("\n".join(body_parts))
        f.write("\n")

    # DOCX
    docx_path = None
    if DOCX_AVAILABLE:
        docx_path = os.path.join(output_dir, f"{base_name}_ocr_output.docx")
        _build_pdf_docx(page_results, source_filename).save(docx_path)

    return txt_path, docx_path