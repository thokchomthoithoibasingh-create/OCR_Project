"""
output_writer.py
-----------------
Takes raw OCR results and saves to both .txt and .docx files.

CHANGE LOG (line-merging fix)
------------------------------
PaddleOCR's detection model returns one entry per detected *visual* text
line — not one entry per *logical* paragraph/sentence. When a numbered
question's number sits in its own little bounding box, slightly separated
from the question text (common with scanned/photographed pages, uneven
spacing, or a numbering column), that number was previously written out
as its own standalone paragraph in the DOCX/TXT output:

    4.
    What are the different modes of data transfer?

instead of:

    4. What are the different modes of data transfer?

To fix this without a full layout model, we use the only spatial signal
available from OCREngine's output (the 4-corner `box` for each line) to
decide which consecutive lines belong in the same paragraph:

    1. Estimate each line's height from its box.
    2. Measure the vertical gap from the bottom of one line to the top
       of the next.
    3. If that gap is small relative to typical line height, the lines
       are merged into one paragraph (this is what reattaches a wrapped
       line, or a split "4." to its question text).
    4. A line that is short AND looks like a bare list/numbering marker
       (e.g. "4.", "(b)", "iii.", "a)") is *always* merged forward into
       the next line, regardless of gap size — a lone numbering token is
       never emitted as its own paragraph.

This is still a heuristic (see OCREngine._sort_reading_order's own
docstring for the same caveat) — it will not perfectly reconstruct
multi-column layouts — but it removes the most common and jarring
artifact: numbering torn away from its text.

Additionally, the DOCX output no longer includes the "OCR Output Report"
metadata header/banner. That header is processing metadata, not part of
the source document, so embedding it in the DOCX made every output look
like a scan report rather than a clean copy of the document. The TXT
output still includes the metadata header (useful there, since .txt is
treated as a log/record), but the DOCX now contains only the merged
document text.

CHANGE LOG (confidence robustness fix)
---------------------------------------
merge_lines_into_paragraphs previously assumed every result's
"confidence" value was numeric. If an OCR engine ever emits None for a
low-quality/unscored detection, summing confidences for averaging would
raise TypeError and abort the whole run. All confidence reads now go
through _safe_conf(), which coerces None/non-numeric values to 0.0
instead of crashing. This does not change output for normal numeric
confidences.
"""

import os
import re
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# A line counts as a "bare numbering marker" if, after stripping
# whitespace, it matches things like: "4.", "12.", "(b)", "iii)", "a."
# Short + matches this pattern => always glued to the next line.
_NUMBERING_MARKER_RE = re.compile(
    r"^\(?(\d{1,3}|[a-zA-Z]{1,3}|[ivxlcdm]{1,6})\)?[.)]?$"
)


def _is_bare_marker(text: str) -> bool:
    """True if `text` looks like a standalone list/number marker with no
    actual content of its own (e.g. '4.', '(b)', 'iii.')."""
    candidate = text.strip()
    if not candidate or len(candidate) > 6:
        return False
    return bool(_NUMBERING_MARKER_RE.match(candidate))


def _line_height(item: dict) -> float:
    ys = [point[1] for point in item["box"]]
    return max(ys) - min(ys)


def _line_top(item: dict) -> float:
    return min(point[1] for point in item["box"])


def _line_bottom(item: dict) -> float:
    return max(point[1] for point in item["box"])


def _safe_conf(value) -> float:
    """Coerce a possibly-missing/None confidence value to a float so
    averaging never raises TypeError on low-quality detections."""
    return float(value) if isinstance(value, (int, float)) else 0.0


def merge_lines_into_paragraphs(results: list, gap_ratio: float = 0.6) -> list:
    """
    Merge consecutive OCR line-entries (already in reading order) into
    logical paragraphs.

    Parameters
    ----------
    results : list of dict
        OCR line results, each with 'text', 'confidence', 'box', already
        sorted into reading order (see OCREngine._sort_reading_order).
    gap_ratio : float
        A line is merged with the previous one if the vertical gap
        between them is less than `gap_ratio * average_line_height` of
        the two lines involved. Tuned conservatively (0.6) so genuinely
        separate paragraphs/questions stay separate, while wrapped lines
        and split numbering still merge.

    Returns
    -------
    list of dict
        Each dict has:
            'text'       : merged paragraph text
            'confidence' : average confidence across merged lines
    """
    if not results:
        return []

    paragraphs = []
    current_texts = [results[0]["text"]]
    current_confidences = [_safe_conf(results[0]["confidence"])]
    current_ref = results[0]

    for item in results[1:]:
        prev_is_marker = _is_bare_marker(current_texts[-1])

        gap = _line_top(item) - _line_bottom(current_ref)
        avg_height = (_line_height(current_ref) + _line_height(item)) / 2 or 1
        close_enough = gap <= gap_ratio * avg_height

        if prev_is_marker or close_enough:
            # Glue this line onto the current paragraph. A bare marker
            # ("4.") gets joined with no space-doubling weirdness; normal
            # continuations get a single space.
            if prev_is_marker:
                current_texts[-1] = f"{current_texts[-1]} {item['text']}".strip()
            else:
                current_texts.append(item["text"])
            current_confidences.append(_safe_conf(item["confidence"]))
            current_ref = item
        else:
            paragraphs.append({
                "text": " ".join(current_texts).strip(),
                "confidence": round(sum(current_confidences) / len(current_confidences), 4),
            })
            current_texts = [item["text"]]
            current_confidences = [_safe_conf(item["confidence"])]
            current_ref = item

    paragraphs.append({
        "text": " ".join(current_texts).strip(),
        "confidence": round(sum(current_confidences) / len(current_confidences), 4),
    })

    return paragraphs


def format_results_for_display(results: list, show_confidence: bool = True) -> str:
    if not results:
        return "[No text detected]"
    lines = []
    for item in results:
        if show_confidence:
            conf_pct = _safe_conf(item["confidence"]) * 100
            lines.append(f"{item['text']}    [confidence: {conf_pct:.2f}%]")
        else:
            lines.append(item["text"])
    return "\n".join(lines)


def format_plain_text(results: list) -> str:
    if not results:
        return ""
    merged = merge_lines_into_paragraphs(results)
    return "\n".join(item["text"] for item in merged)


def calculate_average_confidence(results: list) -> float:
    if not results:
        return 0.0
    total = sum(_safe_conf(item["confidence"]) for item in results)
    return round((total / len(results)) * 100, 2)


def _build_image_docx(results: list, source_filename: str):
    """
    Build a clean DOCX containing only the merged document text — no
    OCR metadata/report header. The DOCX is meant to look like the
    source document, not a processing log.
    """
    doc = Document()
    merged = merge_lines_into_paragraphs(results)
    for item in merged:
        doc.add_paragraph(item["text"])
    return doc


def _build_pdf_docx(page_results: list, source_filename: str):
    """
    Build a clean multi-page DOCX: merged paragraph text per page, with
    a page break between pages and no OCR metadata/report header.
    """
    from docx.enum.text import WD_BREAK

    doc = Document()
    for page_num, results in enumerate(page_results, start=1):
        merged = merge_lines_into_paragraphs(results)
        if merged:
            for item in merged:
                doc.add_paragraph(item["text"])
        else:
            doc.add_paragraph("[No text detected on this page]")

        if page_num < len(page_results):
            doc.add_page_break()
    return doc


def save_image_ocr_output(results: list, source_filename: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(source_filename))[0]

    # TXT — keeps the metadata header; this file is treated as a log/record.
    txt_path = os.path.join(output_dir, f"{base_name}_ocr_output.txt")
    avg_conf = calculate_average_confidence(results)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    header = (
        f"OCR Output Report\n"
        f"{'=' * 50}\n"
        f"Source File       : {os.path.basename(source_filename)}\n"
        f"Processed On      : {timestamp}\n"
        f"Lines Detected    : {len(results)}\n"
        f"Average Confidence: {avg_conf}%\n"
        f"{'=' * 50}\n\n"
    )
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(header)
        f.write(format_plain_text(results))
        f.write("\n")

    # DOCX — clean document text only, no metadata header.
    docx_path = None
    if DOCX_AVAILABLE:
        docx_path = os.path.join(output_dir, f"{base_name}_ocr_output.docx")
        _build_image_docx(results, source_filename).save(docx_path)

    return txt_path, docx_path


def save_pdf_ocr_output(page_results: list, source_filename: str, output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    base_name = os.path.splitext(os.path.basename(source_filename))[0]

    # TXT — keeps the metadata header; this file is treated as a log/record.
    txt_path = os.path.join(output_dir, f"{base_name}_ocr_output.txt")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    total_lines = sum(len(p) for p in page_results)
    all_conf = [_safe_conf(item["confidence"]) for page in page_results for item in page]
    avg_conf = round((sum(all_conf) / len(all_conf)) * 100, 2) if all_conf else 0.0
    header = (
        f"OCR Output Report (Multi-Page PDF)\n"
        f"{'=' * 50}\n"
        f"Source File       : {os.path.basename(source_filename)}\n"
        f"Processed On      : {timestamp}\n"
        f"Total Pages       : {len(page_results)}\n"
        f"Total Lines       : {total_lines}\n"
        f"Average Confidence: {avg_conf}%\n"
        f"{'=' * 50}\n\n"
    )
    body_parts = []
    for page_num, results in enumerate(page_results, start=1):
        page_avg = calculate_average_confidence(results)
        page_header = f"\n--- Page {page_num} (Lines: {len(results)}, Avg Confidence: {page_avg}%) ---\n"
        page_text = format_plain_text(results) if results else "[No text detected on this page]"
        body_parts.append(page_header + page_text)
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(header)
        f.write("\n".join(body_parts))
        f.write("\n")

    # DOCX — clean document text only, no metadata header, page breaks
    # between source pages.
    docx_path = None
    if DOCX_AVAILABLE:
        docx_path = os.path.join(output_dir, f"{base_name}_ocr_output.docx")
        _build_pdf_docx(page_results, source_filename).save(docx_path)

    return txt_path, docx_path